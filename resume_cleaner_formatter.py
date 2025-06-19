from docx import Document
from docx.shared import Pt

def create_clean_resume():
    doc = Document()

    doc.add_heading("Ron Burgos", level=0)
    doc.add_paragraph("Construction Project Manager")
    doc.add_paragraph("Wichita, KS • (620) 200-7642 • burgosron09@yahoo.com • linkedin.com/in/ronburgos")

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Performance-driven Construction Project Manager with over 25 years of experience managing large-scale projects "
        "in HVAC-R, plumbing, electrical, roofing, and infrastructure remodeling. Demonstrated success in completing complex projects "
        "on time and under budget while ensuring safety, quality control, and regulatory compliance. Skilled in subcontractor management, "
        "risk mitigation, and energy efficiency systems across institutional, industrial, and government facilities."
    )

    doc.add_heading("Core Competencies", level=1)
    doc.add_paragraph(
        "Construction Project Management • MEP Systems • OSHA & Code Compliance • Energy Retrofits • Budgeting & Estimating • "
        "Risk Mitigation • Field Supervision • Scheduling • Quality Assurance • Vendor Negotiation"
    )

    doc.add_heading("Professional Experience", level=1)

    experience_data = [
        ("Construction Site Manager", "Dowson Global", "Sacramento, CA", "Jul 2024 – Present", [
            "Lead MEP-heavy construction projects from design through execution across multiple sites.",
            "Coordinate cross-functional teams, conduct risk assessments, and oversee code compliance.",
            "Drive project delivery on tight timelines and budgets, while managing stakeholder expectations."
        ]),
        ("Senior Construction Site Manager", "Iconergy", "Denver, CO", "Feb 2019 – Jul 2024", [
            "Directed energy conservation upgrades for schools, prisons, and hospitals nationwide.",
            "Managed retrofits totaling $30M+, including HVAC, LED lighting, and water systems.",
            "Led 16+ subcontractors and ensured regulatory and safety compliance."
        ]),
        ("Estimator / HVAC Coordinator", "Johnson Controls", "Wichita, KS", "Feb 2018 – Aug 2018", [
            "Estimated and coordinated HVAC installations, boiler and rooftop unit projects.",
            "Managed client communications, procurement, and preventive maintenance contracts."
        ])
    ]

    for title, company, location, dates, bullets in experience_data:
        job_header = doc.add_paragraph()
        job_header.add_run(f"{title} – {company}").bold = True
        job_header.add_run(f" ({location})")
        doc.add_paragraph(f"{dates}")
        for b in bullets:
            doc.add_paragraph(f"• {b}", style='List Bullet')

    doc.add_heading("Education", level=1)
    doc.add_paragraph("Mechanical Technology Certificate – Salina Technical College, Salina, KS")
    doc.add_paragraph("Small Business Mgmt. – Cloud County Community College, Abilene, KS")
    doc.add_paragraph("Class “A” General Contractors License")
    doc.add_paragraph("Master Contractors Mechanical License")
    doc.add_paragraph("Journeyman Plumbers License")
    doc.add_paragraph("Industrial Electrical & HVAC Certified – NCCER")
    doc.add_paragraph("OSHA 10 & 30 Certificate")
    doc.add_paragraph("Pursuing Project Management Certification – University of Phoenix")

    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph("AutoCAD • Procore • Bluebeam Revu • Microsoft Project • Excel • PlanGrid • Smartsheet")

    doc.add_heading("Affiliations", level=1)
    doc.add_paragraph("Member – Associated General Contractors (AGC)")
    doc.add_paragraph("Member – Project Management Institute (PMI)")

    return doc
